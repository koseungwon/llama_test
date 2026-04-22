import subprocess
import tempfile
from pathlib import Path

import pdfplumber
import openpyxl
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".hwp"}


def load_document(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    loaders = {
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".doc": _load_docx,
        ".xlsx": _load_excel,
        ".xls": _load_excel,
        ".hwp": _load_hwp,
    }
    loader = loaders.get(ext)
    if not loader:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")
    return loader(file_path)


def _load_pdf(file_path: str) -> str:
    texts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
    return "\n".join(texts)


def _load_docx(file_path: str) -> str:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 표 내용도 추출
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))
    return "\n".join(paragraphs)


def _load_excel(file_path: str) -> str:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    texts = []
    for sheet in wb.worksheets:
        texts.append(f"[시트: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell) for cell in row if cell is not None]
            if row_values:
                texts.append(" | ".join(row_values))
    return "\n".join(texts)


def _load_hwp(file_path: str) -> str:
    # pyhwp의 hwp5txt CLI를 사용해 텍스트 추출
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        tmp_path = tmp.name

    result = subprocess.run(
        ["hwp5txt", "--output", tmp_path, file_path],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"HWP 변환 실패: {result.stderr}")

    with open(tmp_path, encoding="utf-8", errors="ignore") as f:
        return f.read()
